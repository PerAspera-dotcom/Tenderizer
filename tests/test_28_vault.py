"""Step 28 — Vault ingest slice: chunk/parse (real), embed/extract (mocked —
no live model download or Anthropic call in the test suite), store round-
trips, and the /api/vault/* endpoints (background processing monkeypatched
to a fast fake, same reasoning as not calling the real pipeline in tests).
"""
import asyncio
import io

from fastapi import BackgroundTasks, UploadFile

import store
import vault
import api
from conftest import TEST_TENANT_ID


# ── chunk_text ────────────────────────────────────────────────────────────────

def test_chunk_text_windows_with_overlap():
    text = " ".join(f"w{i}" for i in range(500))
    chunks = vault.chunk_text(text)
    assert len(chunks) == 2
    assert chunks[0].split()[0] == "w0"
    assert chunks[0].split()[-1] == "w399"
    # step = 400-50 = 350, so chunk 2 starts at word 350
    assert chunks[1].split()[0] == "w350"


def test_chunk_text_drops_short_chunks():
    assert vault.chunk_text("short text") == []


def test_chunk_text_empty_string():
    assert vault.chunk_text("") == []


# ── parse_document ────────────────────────────────────────────────────────────

def test_parse_document_unsupported_type_returns_none(tmp_path):
    p = tmp_path / "image.png"
    p.write_bytes(b"\x89PNG\r\n")
    assert vault.parse_document(str(p), "image/png") is None


def test_parse_document_docx(tmp_path):
    from docx import Document
    p = tmp_path / "spec.docx"
    d = Document()
    d.add_paragraph("Material: 600D polyester")
    d.add_paragraph("Water column: 3000 mm")
    d.save(str(p))
    text = vault.parse_document(str(p), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    assert "600D polyester" in text
    assert "3000 mm" in text


# ── ingest_and_embed (embedding model + Chroma mocked) ───────────────────────

class _FakeModel:
    def encode(self, texts):
        class _Arr(list):
            def tolist(self):
                return [[0.0, 0.1, 0.2] for _ in texts]
        return _Arr()


class _FakeCollection:
    def __init__(self):
        self.upserted = None

    def upsert(self, ids, documents, embeddings, metadatas):
        self.upserted = {"ids": ids, "documents": documents,
                          "embeddings": embeddings, "metadatas": metadatas}


def test_ingest_and_embed_upserts_chunks(tmp_path, monkeypatch):
    fake_collection = _FakeCollection()
    monkeypatch.setattr(vault, "_embedding_model", lambda: _FakeModel())
    monkeypatch.setattr(vault, "_chroma_collection", lambda tenant_id: fake_collection)

    p = tmp_path / "spec.txt"
    monkeypatch.setattr(vault, "_is_pdf", lambda path, ct: False)
    monkeypatch.setattr(vault, "_is_docx", lambda path, ct: False)
    monkeypatch.setattr(vault, "parse_document", lambda path, ct: " ".join(f"w{i}" for i in range(500)))

    n = vault.ingest_and_embed(TEST_TENANT_ID, doc_id=1, path=str(p), content_type="text/plain")
    assert n == 2
    assert fake_collection.upserted["ids"] == ["doc1_chunk0", "doc1_chunk1"]


def test_ingest_and_embed_no_text_stores_nothing(tmp_path, monkeypatch):
    monkeypatch.setattr(vault, "parse_document", lambda path, ct: "")
    n = vault.ingest_and_embed(TEST_TENANT_ID, doc_id=1, path="whatever", content_type="text/plain")
    assert n == 0


# ── extract_metadata ──────────────────────────────────────────────────────────

def test_extract_metadata_no_api_key_returns_none(monkeypatch):
    monkeypatch.delenv("ANTHROPIC_API_KEY", raising=False)
    assert vault.extract_metadata("doc.pdf", "application/pdf") is None


def test_extract_metadata_non_pdf_returns_none(monkeypatch):
    monkeypatch.setenv("ANTHROPIC_API_KEY", "fake-key-for-test")
    assert vault.extract_metadata("doc.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document") is None


def test_extract_metadata_parses_claude_json_response(monkeypatch):
    monkeypatch.setenv("ANTHROPIC_API_KEY", "fake-key-for-test")
    monkeypatch.setattr(vault, "_pdf_first_pages_as_images", lambda path: ["base64imagedata"])

    class _FakeContentBlock:
        text = ('{"doc_type": "Datasheet", "metadata": {"Material": "600D PES"}, '
                '"cpv_codes": ["39522530"], "confidence": 0.9}')

    class _FakeMessage:
        content = [_FakeContentBlock()]

    class _FakeMessages:
        def create(self, **kwargs):
            return _FakeMessage()

    class _FakeClient:
        messages = _FakeMessages()

    import anthropic
    monkeypatch.setattr(anthropic, "Anthropic", lambda: _FakeClient())

    result = vault.extract_metadata("doc.pdf", "application/pdf")
    doc_type, metadata, cpv_codes, confidence = result
    assert doc_type == "Datasheet"
    assert metadata == {"Material": "600D PES"}
    assert cpv_codes == ["39522530"]
    assert confidence == 0.9


def test_parse_metadata_response_handles_code_fence():
    text = '```json\n{"doc_type": "Certificate", "metadata": {}, "cpv_codes": [], "confidence": 0.2}\n```'
    parsed = vault._parse_metadata_response(text)
    assert parsed["doc_type"] == "Certificate"


def test_parse_metadata_response_invalid_json_returns_none():
    assert vault._parse_metadata_response("not json at all") is None


def test_parse_metadata_response_missing_metadata_key_returns_none():
    assert vault._parse_metadata_response('{"doc_type": "Other"}') is None


# ── store round-trip ──────────────────────────────────────────────────────────

def test_add_and_list_vault_document_defaults_to_processing(tmp_path):
    conn = store.init_db(str(tmp_path / "t.db"))
    doc_id = store.add_vault_document(conn, TEST_TENANT_ID, "spec.pdf", "application/pdf",
                                       1234, "/fake/path")
    docs = store.list_vault_documents(conn, TEST_TENANT_ID)
    assert len(docs) == 1
    assert docs[0]["id"] == doc_id
    assert docs[0]["status"] == "processing"
    assert docs[0]["metadata"] == {}
    assert docs[0]["cpv_codes"] == []
    assert docs[0]["confidence"] is None


def test_list_vault_documents_search_by_filename(tmp_path):
    conn = store.init_db(str(tmp_path / "t.db"))
    store.add_vault_document(conn, TEST_TENANT_ID, "tent-fabric-spec.pdf", "application/pdf", 1, "/p1")
    store.add_vault_document(conn, TEST_TENANT_ID, "iso-certificate.pdf", "application/pdf", 1, "/p2")
    assert len(store.list_vault_documents(conn, TEST_TENANT_ID, q="fabric")) == 1
    assert len(store.list_vault_documents(conn, TEST_TENANT_ID, q="cert")) == 1
    assert len(store.list_vault_documents(conn, TEST_TENANT_ID)) == 2


def test_update_vault_document_metadata_flips_to_indexed(tmp_path):
    conn = store.init_db(str(tmp_path / "t.db"))
    doc_id = store.add_vault_document(conn, TEST_TENANT_ID, "spec.pdf", "application/pdf", 1, "/p1")
    store.update_vault_document_metadata(
        conn, TEST_TENANT_ID, doc_id, doc_type="Datasheet",
        metadata={"Material": "600D PES"}, cpv_codes=["39522530"],
        confidence=0.9, fields_extracted=1, status="indexed")
    doc = store.list_vault_documents(conn, TEST_TENANT_ID)[0]
    assert doc["status"] == "indexed"
    assert doc["doc_type"] == "Datasheet"
    assert doc["metadata"] == {"Material": "600D PES"}
    assert doc["confidence"] == 0.9


def test_get_vault_document_is_tenant_scoped(tmp_path):
    conn = store.init_db(str(tmp_path / "t.db"))
    doc_id = store.add_vault_document(conn, TEST_TENANT_ID, "spec.pdf", "application/pdf", 1, "/p1")
    assert store.get_vault_document(conn, TEST_TENANT_ID, doc_id)["filename"] == "spec.pdf"
    assert store.get_vault_document(conn, 999, doc_id) is None


# ── API endpoints ─────────────────────────────────────────────────────────────

def _seed(tmp_path, monkeypatch):
    db_path = str(tmp_path / "t.db")
    monkeypatch.setattr(api, "DB_PATH", db_path)
    monkeypatch.setattr(api, "VAULT_UPLOAD_DIR", tmp_path / "vault_uploads")
    return store.init_db(db_path)


def _file(name="spec.pdf", content=b"%PDF-1.4 fake pdf content"):
    return UploadFile(file=io.BytesIO(content), filename=name)


def test_ingest_vault_document_returns_processing_row(tmp_path, monkeypatch):
    _seed(tmp_path, monkeypatch)
    background = BackgroundTasks()
    result = asyncio.run(api.ingest_vault_document(background, file=_file(), tenant_id=TEST_TENANT_ID))
    assert result["filename"] == "spec.pdf"
    assert result["status"] == "processing"
    assert len(background.tasks) == 1  # queued, not run inline


def test_get_vault_docs_shape_matches_frontend_contract(tmp_path, monkeypatch):
    conn = _seed(tmp_path, monkeypatch)
    store.add_vault_document(conn, TEST_TENANT_ID, "a.pdf", "application/pdf", 1, "/p1")
    store.update_vault_document_metadata(conn, TEST_TENANT_ID, 1, doc_type="Datasheet",
                                          metadata={"Material": "PES"}, cpv_codes=[],
                                          confidence=0.8, fields_extracted=1, status="indexed")
    store.add_vault_document(conn, TEST_TENANT_ID, "b.pdf", "application/pdf", 1, "/p2")

    body = api.get_vault_docs(q=None, tenant_id=TEST_TENANT_ID)
    assert body["total"] == 2
    assert body["processing"] == 1
    assert {d["filename"] for d in body["results"]} == {"a.pdf", "b.pdf"}


def test_get_vault_docs_search_filters_results(tmp_path, monkeypatch):
    conn = _seed(tmp_path, monkeypatch)
    store.add_vault_document(conn, TEST_TENANT_ID, "tent-fabric.pdf", "application/pdf", 1, "/p1")
    store.add_vault_document(conn, TEST_TENANT_ID, "certificate.pdf", "application/pdf", 1, "/p2")

    body = api.get_vault_docs(q="fabric", tenant_id=TEST_TENANT_ID)
    assert body["total"] == 1
    assert body["results"][0]["filename"] == "tent-fabric.pdf"


def test_run_vault_processing_updates_store(tmp_path, monkeypatch):
    conn = _seed(tmp_path, monkeypatch)
    doc_id = store.add_vault_document(conn, TEST_TENANT_ID, "spec.pdf", "application/pdf", 1, "/p1")
    monkeypatch.setattr(vault, "process_upload", lambda tenant_id, doc_id, path, content_type: {
        "doc_type": "Datasheet", "metadata": {"Material": "PES"}, "cpv_codes": ["39522530"],
        "confidence": 0.9, "fields_extracted": 1, "status": "indexed"})

    api._run_vault_processing(TEST_TENANT_ID, doc_id, "/p1", "application/pdf")

    doc = store.list_vault_documents(api._db(), TEST_TENANT_ID)[0]
    assert doc["status"] == "indexed"
    assert doc["doc_type"] == "Datasheet"
