"""Composer — per-tender proposal-drafting pipeline.

Wraps the standalone `proposal_tool/*.py` scripts (a working, human-run
Python pipeline) as per-tender, importable functions — following the same
pattern `vault.py` already established for Vault: per-tenant/per-tender
paths, plain functions with explicit params, `ANTHROPIC_API_KEY` from the
environment (never hardcoded), additive (not destructive) Chroma upserts.

Reuses `vault.chunk_text`/`vault.parse_document`/`vault._embedding_model`
directly rather than reimplementing identical parsing/chunking logic.

Two things `proposal_tool/generate.py` does are kept deliberately separate
here:
  - The SOW-extracted requirement list (`extract_requirements`) drives the
    Proposal Review / Gaps Report screens and the generated `.docx` — this
    is new logic; the mockup's per-requirement confidence/source-ref chip
    has no equivalent in the original pipeline, which reads its 35
    requirements verbatim from `compliance_matrix.xlsx` rows instead.
  - `fill_compliance_matrix` ports that original per-matrix-row/per-tent-size
    logic (`COL_*`, `TENT_SIZES`, the YES_NO/CROSS_REFERENCE/REMARK prompt
    protocol) nearly unchanged, as an independent, optional deliverable —
    not reconciled against the AI-extracted requirements above.
"""
import base64
import json
import os
from datetime import datetime

import chromadb
import openpyxl
import pandas as pd
import pdfplumber
from docx import Document
from docx.shared import RGBColor
from openpyxl.styles import Font as XLFont, PatternFill

import vault

CHROMA_ROOT = "data/composer_chroma"
COLLECTION_NAME = "composer_docs"
OUTPUT_ROOT = "data/composer_output"
CLAUDE_MODEL = vault.CLAUDE_MODEL

GOOD_SIMILARITY = 0.35
PARTIAL_SIMILARITY = 0.20
TOP_K = 5
MAX_ENRICH_PAGES = 20
ENRICH_DPI = 150
IMAGE_HEAVY_CHAR_THRESHOLD = 40  # avg extracted chars/page below this -> likely scanned/image-only

ROLE_MAP = {
    "sow_": "sow",
    "tech_": "tech",
    "background_": "background",
    "parta_": "parta",
    "example_": "example",
}


def get_role(filename):
    """Case-insensitive filename-prefix role detection — ported verbatim
    from proposal_tool/ingest.py's ROLE_MAP/get_role. "unknown" for anything
    unrecognised (still ingested, just untagged).
    """
    lower = filename.lower()
    for prefix, role in ROLE_MAP.items():
        if lower.startswith(prefix):
            return role
    return "unknown"


_PARTA_START_MARKERS = [
    "part a - technical proposal/capability and qualification form",
    "part a – technical proposal/capability and qualification form",
    "part a: technical proposal/capability and qualification form",
    "part a - technical proposal / capability and qualification form",
    "part a – technical proposal / capability and qualification form",
    "part a - technical proposal capability and qualification form",
    "part a: technical proposal capability and qualification form",
    "part a tp capability and qualification",
    "part a: tp capability and qualification",
    "part a - tp capability and qualification",
    "part a–tp capability and qualification",
]
_PARTA_END_MARKERS = ["part b", "part c", "part d", "part e", "annex", "appendix",
                       "section ii", "section 2"]


def extract_parta_section(text):
    """Clip to the Part A capability/qualification section — ported from
    proposal_tool/ingest.py's extract_parta_section, minus its print()s.
    Falls back to the full text (with a low confidence signal left to the
    caller) if no recognised header is found, same as the original.
    """
    lower = text.lower()
    start = -1
    for marker in _PARTA_START_MARKERS:
        idx = lower.find(marker)
        if idx != -1:
            start = idx
            break
    if start == -1:
        return text

    end = len(text)
    for marker in _PARTA_END_MARKERS:
        idx = lower.find(marker, start + 200)
        if idx != -1 and idx < end:
            end = idx
    return text[start:end]


def _chroma_collection(tenant_id, pub_number):
    client = chromadb.PersistentClient(path=f"{CHROMA_ROOT}/{tenant_id}/{pub_number}")
    return client.get_or_create_collection(COLLECTION_NAME, metadata={"hnsw:space": "cosine"})


def _pdf_pages_text(path):
    pages = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            pages.append(page.extract_text() or "")
    return pages


def detect_image_heavy(path, content_type):
    """True if this looks like a scanned/image-only PDF (little to no
    extractable text per page) — the trigger for the "run datasheet
    enrichment" prompt in the Ingest & Config UI. Never true for non-PDFs.
    """
    is_pdf = content_type == "application/pdf" or path.lower().endswith(".pdf")
    if not is_pdf:
        return False
    pages = _pdf_pages_text(path)
    if not pages:
        return False
    avg_chars = sum(len(p) for p in pages) / len(pages)
    return avg_chars < IMAGE_HEAVY_CHAR_THRESHOLD


def ingest_document(tenant_id, pub_number, doc_id, path, content_type, role):
    """Parse -> (Part A clip, if applicable) -> chunk -> embed -> upsert into
    this tender's Chroma collection. Additive (get_or_create + upsert), never
    the original script's delete-and-rebuild-everything, so ingesting one
    document never destroys another tender's or another document's chunks.
    Returns the number of chunks stored (0 for unsupported/empty files —
    not an error).
    """
    text = vault.parse_document(path, content_type)
    if not text or not text.strip():
        return 0
    if role == "parta":
        text = extract_parta_section(text)
    chunks = vault.chunk_text(text)
    if not chunks:
        return 0
    embeddings = vault._embedding_model().encode(chunks).tolist()
    ids = [f"doc{doc_id}_chunk{i}" for i in range(len(chunks))]
    metadatas = [{"source": os.path.basename(path), "doc_id": doc_id, "role": role} for _ in chunks]
    collection = _chroma_collection(tenant_id, pub_number)
    collection.upsert(ids=ids, documents=chunks, embeddings=embeddings, metadatas=metadatas)
    return len(chunks)


_ENRICH_PROMPT = """You are extracting technical information from a product datasheet.
Extract ALL of the following that are present on this page:

- Product name and model number
- Dimensions (length, width, height, weight — include all units)
- Materials and fabric specifications (include technical codes, weights, coatings)
- Performance specifications (waterproofing ratings, wind resistance, temperature range)
- Certifications and standards compliance (ISO, EN, NATO STANAG, etc.)
- Component list (poles, pegs, guylines, bags, accessories)
- Colour options and camouflage patterns
- Packing dimensions and packed weight
- Manufacturing country and facility information
- Test results and test standards referenced
- Any other technical specifications shown

Format your output as structured text with clear headings.
Be precise — copy exact numbers, codes and specifications as shown.
If a value is in a table or diagram, extract it.
If a page contains no technical information (e.g. cover page, blank page), say: NO TECHNICAL CONTENT.
"""


def enrich_datasheet(path):
    """Claude-Vision per-page spec extraction for an image-heavy PDF — ported
    from proposal_tool/enrich_datasheets.py's enrich_document(), minus its
    file-caching side effect (the caller re-ingests the returned text
    directly rather than writing a sibling _extracted.txt). Returns the
    concatenated extracted text (empty string if no page had content), or
    None if ANTHROPIC_API_KEY isn't configured.
    """
    if not os.getenv("ANTHROPIC_API_KEY"):
        return None
    import anthropic
    import fitz  # PyMuPDF

    client = anthropic.Anthropic()
    doc = fitz.open(path)
    try:
        page_count = min(len(doc), MAX_ENRICH_PAGES)
        extracted_pages = []
        for i in range(page_count):
            page = doc[i]
            mat = fitz.Matrix(ENRICH_DPI / 72, ENRICH_DPI / 72)
            pix = page.get_pixmap(matrix=mat, colorspace=fitz.csRGB)
            b64 = base64.standard_b64encode(pix.tobytes("png")).decode("utf-8")
            message = client.messages.create(
                model=CLAUDE_MODEL, max_tokens=2000,
                messages=[{"role": "user", "content": [
                    {"type": "image", "source": {"type": "base64", "media_type": "image/png", "data": b64}},
                    {"type": "text", "text": _ENRICH_PROMPT},
                ]}],
            )
            text = message.content[0].text
            if "NO TECHNICAL CONTENT" not in text.upper():
                extracted_pages.append(text.strip())
        return "\n\n".join(extracted_pages)
    finally:
        doc.close()


_REQUIREMENTS_PROMPT = """You are extracting distinct, actionable requirements from a tender's \
Statement of Work (SOW) and/or capability/qualification form, for a company preparing a \
technical proposal response.

The source text below is tagged with [<document> · p.<page>] markers before each page's content.

Extract every distinct technical, functional, or administrative requirement the supplier must \
satisfy. Combine near-duplicate mentions of the same requirement into a single entry. Skip pure \
boilerplate (cover pages, tables of contents, generic legal disclaimers) unless a requirement is \
embedded in it.

Return ONLY a JSON array (no markdown fences, no other text), each item shaped exactly:
{{
  "title": "<short requirement title, under 12 words>",
  "extracted": "<the requirement, quoted or closely paraphrased from the source text>",
  "source": "<document name, plus a section number if visible in the text (e.g. '§4.2'), plus \
the page it was found on, e.g. 'sow_tender.pdf §4.2 · p.12'>",
  "confidence": <your own confidence this is a genuine, distinct, actionable requirement, 0.0-1.0>
}}

SOURCE TEXT:
{text}
"""

# A very long SOW (100+ pages) would need chunking + cross-call dedup to stay
# under context limits — out of scope for this pass; typical tender SOWs
# (tens of pages) fit comfortably in one call.
_MAX_REQUIREMENTS_CONTEXT_CHARS = 180_000


def extract_requirements(documents):
    """documents: [{"filename": str, "role": "sow"|"parta", "pages": [str, ...]}]
    (parta docs should already be clipped via extract_parta_section before
    being passed in). Returns [{"title", "extracted", "source", "confidence"}],
    or [] if ANTHROPIC_API_KEY isn't configured or nothing parseable came
    back — same graceful-skip convention as vault.extract_metadata.
    """
    if not os.getenv("ANTHROPIC_API_KEY"):
        return []
    blocks = []
    for doc in documents:
        for i, page_text in enumerate(doc.get("pages") or [], 1):
            if page_text and page_text.strip():
                blocks.append(f"[{doc['filename']} · p.{i}]\n{page_text.strip()}")
    if not blocks:
        return []
    text = "\n\n".join(blocks)[:_MAX_REQUIREMENTS_CONTEXT_CHARS]

    import anthropic
    client = anthropic.Anthropic()
    message = client.messages.create(
        model=CLAUDE_MODEL, max_tokens=4096,
        messages=[{"role": "user", "content": _REQUIREMENTS_PROMPT.format(text=text)}],
    )
    return _parse_requirements_response(message.content[0].text)


def _parse_requirements_response(text):
    cleaned = text.strip()
    if cleaned.startswith("```"):
        cleaned = cleaned.strip("`")
        if cleaned.lower().startswith("json"):
            cleaned = cleaned[4:]
    try:
        data = json.loads(cleaned)
    except (TypeError, ValueError):
        return []
    if not isinstance(data, list):
        return []
    out = []
    for item in data:
        if not isinstance(item, dict) or not item.get("extracted"):
            continue
        confidence = item.get("confidence")
        out.append({
            "title": item.get("title") or item["extracted"][:80],
            "extracted": item["extracted"],
            "source": item.get("source") or "",
            "confidence": confidence if isinstance(confidence, (int, float)) else None,
        })
    return out


def retrieve_evidence(tenant_id, pub_number, query, roles, top_k=TOP_K):
    """Chroma query restricted to `roles`, cosine similarity conversion
    identical to proposal_tool/generate.py's retrieve_by_role. [] if this
    tender's collection has nothing indexed yet.
    """
    collection = _chroma_collection(tenant_id, pub_number)
    if collection.count() == 0:
        return []
    embedding = vault._embedding_model().encode([query]).tolist()
    results = collection.query(
        query_embeddings=embedding, n_results=min(top_k * 4, 40),
        include=["documents", "metadatas", "distances"], where={"role": {"$in": roles}},
    )
    chunks = []
    for doc, meta, dist in zip(results["documents"][0], results["metadatas"][0], results["distances"][0]):
        similarity = round(1 - (dist / 2), 3)
        chunks.append({"text": doc, "source": meta["source"], "role": meta["role"], "similarity": similarity})
    return chunks[:top_k]


def _gap_status(chunks):
    if not chunks:
        return "completed"
    best = max(c["similarity"] for c in chunks)
    if best >= GOOD_SIMILARITY:
        return "complete"
    if best >= PARTIAL_SIMILARITY:
        return "linked"
    return "completed"


def _evidence_block(chunks):
    if not chunks:
        return "No relevant technical documentation found."
    return "".join(
        f"\n[Source {i}: {c['source']} | relevance: {c['similarity']}]\n{c['text']}\n"
        for i, c in enumerate(chunks, 1)
    )


_RESPONSE_PROMPT = """You are writing a section of a Technical Proposal responding to a \
government tender.

REQUIREMENT:
{requirement}

RETRIEVED TECHNICAL EVIDENCE:
{evidence}
{style_block}
Write a 150-300 word formal proposal response. Structure: opening claim of compliance -> \
elaboration of how/what -> evidence from the retrieved documentation. Ground every claim in the \
retrieved evidence — do not invent specifications or certifications. Return only the prose, no \
headings, no preamble.
"""


def generate_response(requirement_text, evidence_chunks, style_guide=None):
    style_block = f"\nWRITING STYLE GUIDE — follow this precisely:\n{style_guide}\n" if style_guide else ""
    prompt = _RESPONSE_PROMPT.format(
        requirement=requirement_text, evidence=_evidence_block(evidence_chunks), style_block=style_block)
    import anthropic
    client = anthropic.Anthropic()
    message = client.messages.create(model=CLAUDE_MODEL, max_tokens=1000,
                                      messages=[{"role": "user", "content": prompt}])
    return message.content[0].text.strip()


_REFINE_PROMPT = """You wrote this proposal response:
{current}

The reviewer gave this feedback:
{feedback}

Additional retrieved evidence:
{evidence}

Rewrite the response incorporating the feedback. Keep it 150-250 words. Return only the prose."""


def refine_section(requirement_text, current_response, feedback, evidence_chunks):
    """Section-scoped regenerate — the real, working version of
    proposal_tool/refine.py's logic (that script assigns a bare API-key
    string where an anthropic.Anthropic() client is expected, so
    `claude.messages.create(...)` there would raise; this is a rewrite, not
    a port, using the env-var client like everywhere else in this module).
    """
    prompt = _REFINE_PROMPT.format(current=current_response, feedback=feedback,
                                    evidence=_evidence_block(evidence_chunks))
    import anthropic
    client = anthropic.Anthropic()
    message = client.messages.create(model=CLAUDE_MODEL, max_tokens=800,
                                      messages=[{"role": "user", "content": prompt}])
    return message.content[0].text.strip()


def run_generate(tenant_id, pub_number, requirements, style_guide=None):
    """requirements: [{id, title, extracted}] (validated requirements only —
    the caller enforces the validation gate). For each: retrieve tech
    evidence, derive gap_status from similarity, and generate a response
    unless there's no evidence at all. Returns
    [{id, gap_status, similarity, response_text, citations}] in the same
    order given — the caller persists these via
    store.update_composer_requirement_result.
    """
    results = []
    for req in requirements:
        query = f"{req['title']} {req['extracted']}"
        tech_chunks = retrieve_evidence(tenant_id, pub_number, query, roles=["tech"])
        status = _gap_status(tech_chunks)
        best_sim = max((c["similarity"] for c in tech_chunks), default=0.0)
        response_text, citations = None, []
        if status != "completed":
            citations = [{"doc": c["source"], "score": c["similarity"]} for c in tech_chunks]
            # Still worth persisting gap_status/similarity/citations without a
            # key configured (informative for Gaps Report); only the prose
            # generation itself needs Claude.
            if os.getenv("ANTHROPIC_API_KEY"):
                response_text = generate_response(req["extracted"], tech_chunks, style_guide)
        results.append({"id": req["id"], "gap_status": status, "similarity": best_sim,
                         "response_text": response_text, "citations": citations})
    return results


def _add_red_run(paragraph, text):
    run = paragraph.add_run(text)
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    run.italic = True
    return run


def build_proposal_docx(requirements, out_path, background_text=None):
    """requirements: the full persisted shape (store.list_composer_requirements'
    rows) — needs title/extracted/gap_status/response_text/citations. Section
    structure ported from proposal_tool/generate.py, adapted to one
    subsection per SOW-extracted requirement rather than per
    (matrix-row × tent-size).
    """
    doc = Document()
    doc.add_heading("Technical Proposal", 0)

    doc.add_heading("Section 1 — Introduction", level=1)
    intro = doc.add_paragraph()
    intro.add_run("This Technical Proposal is submitted by ")
    _add_red_run(intro, "[Company Name]")
    intro.add_run(" in response to the request for tender issued by ")
    _add_red_run(intro, "[Tendering Company / Authority]")
    intro.add_run(".")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    doc.add_heading("Section 2 — Company Background", level=1)
    if background_text:
        for para in background_text.split("\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
    else:
        missing = doc.add_paragraph()
        _add_red_run(missing, "To be completed — add a background_ prefixed document to the "
                               "docs folder containing company background information.")

    doc.add_heading("Section 3 — Scope of Requirements", level=1)
    doc.add_paragraph("The following sections address each requirement extracted from the "
                       "Statement of Work. All responses are substantiated by referenced "
                       "technical documentation.")

    cited_docs, to_link_docs = set(), set()
    for i, req in enumerate(requirements, 1):
        section_num = f"3.{i}"
        doc.add_heading(f"Section {section_num} — {req['title']}", level=2)
        req_para = doc.add_paragraph()
        req_para.add_run(req["extracted"]).italic = True

        if req["gap_status"] == "completed":
            blank = doc.add_paragraph()
            _add_red_run(blank, "To be completed — no technical documentation found for this "
                                 "requirement. Please provide evidence of compliance and "
                                 "complete this section before submission.")
        else:
            doc.add_paragraph(req["response"] or "")
            if req["gap_status"] == "linked" and req["citations"]:
                doc_name = req["citations"][0]["doc"]
                to_link_docs.add(doc_name)
                link_para = doc.add_paragraph()
                link_para.add_run("Supporting document: ")
                _add_red_run(link_para, f"{doc_name} — to be linked")
            else:
                for c in req["citations"]:
                    cited_docs.add(c["doc"])

    doc.add_heading("Annex A — Referenced Technical Documents", level=1)
    doc.add_paragraph("The following technical documents have been referenced in this proposal "
                       "as evidence of compliance:")
    if cited_docs:
        for name in sorted(cited_docs):
            doc.add_paragraph(name, style="List Bullet")
    else:
        no_docs = doc.add_paragraph()
        _add_red_run(no_docs, "To be completed — no technical documents were cited. Add tech_ "
                               "prefixed documents and re-run.")
    if to_link_docs:
        doc.add_paragraph("")
        doc.add_paragraph("The following documents require formal linking before submission:")
        for name in sorted(to_link_docs):
            link_para = doc.add_paragraph()
            _add_red_run(link_para, f"{name} — to be linked")

    doc.save(out_path)
    return out_path


def build_gaps_report(requirements, out_path):
    """Plaintext gaps_report.txt in the same format proposal_tool/generate.py
    produces (confirmed against a real sample). The structured equivalent is
    just `requirements` itself — the API returns that directly, no
    text-reparsing needed on the way back.
    """
    completed = [r for r in requirements if r["gap_status"] == "completed"]
    linked = [r for r in requirements if r["gap_status"] == "linked"]
    complete = [r for r in requirements if r["gap_status"] == "complete"]
    total = len(requirements)
    gaps_exist = bool(completed or linked)

    with open(out_path, "w", encoding="utf-8") as f:
        f.write("PROPOSAL GAPS REPORT\n")
        f.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
        f.write("=" * 60 + "\n\n")
        f.write(f"SUBMISSION READINESS: {'NOT READY' if gaps_exist else 'READY'}\n\n")
        f.write(f"Total requirements:  {total}\n")
        f.write(f"  Complete (ready):           {len(complete)}\n")
        f.write(f"  To be completed (no docs):  {len(completed)}\n")
        f.write(f"  To be linked (weak docs):   {len(linked)}\n\n")

        if completed:
            f.write("─" * 60 + "\n")
            f.write("TO BE COMPLETED — no technical documentation found:\n")
            f.write("(Add relevant tech_ documents and re-run to resolve)\n\n")
            for r in completed:
                f.write(f"  {r['title']}\n")
                f.write(f"  {r['extracted'][:100]}\n\n")

        if linked:
            f.write("─" * 60 + "\n")
            f.write("TO BE LINKED — document found but needs formal linking:\n\n")
            for r in linked:
                doc_name = r["citations"][0]["doc"] if r["citations"] else "unknown"
                f.write(f"  {r['title']}\n")
                f.write(f"  Document: {doc_name}\n")
                f.write(f"  {r['extracted'][:100]}\n\n")

        f.write("=" * 60 + "\n")
        f.write("Proposal is not submission-ready until all items above are resolved.\n")
    return out_path


# ── Compliance matrix fill — independent of the SOW-extracted requirements ──
# above. Ported near-verbatim from proposal_tool/generate.py: same column
# map (confirmed against the real compliance_matrix.xlsx/matrix_filled.xlsx
# samples), same tent-size loop, same YES_NO/CROSS_REFERENCE/REMARK protocol.
# Two independent consumers of the same per-tender evidence index, not a
# forced mapping between AI-extracted requirements and matrix rows.

COL_NO, COL_REQ, COL_SOW_REF, COL_DOC_GUIDE, COL_VERDELING = 0, 1, 2, 12, 13

TENT_SIZES = [
    {"label": "1 person tent",  "col_yesno": 3,  "col_ref": 4,  "col_remark": 5},
    {"label": "2 person tent",  "col_yesno": 6,  "col_ref": 7,  "col_remark": 8},
    {"label": "3 person tent",  "col_yesno": 9,  "col_ref": 10, "col_remark": 11},
    {"label": "4 person tent",  "col_yesno": 14, "col_ref": 15, "col_remark": 16},
    {"label": "6 person tent",  "col_yesno": 17, "col_ref": 18, "col_remark": 19},
    {"label": "8 person tent",  "col_yesno": 20, "col_ref": 21, "col_remark": 22},
    {"label": "14 person tent", "col_yesno": 23, "col_ref": 24, "col_remark": 25},
]

_MATRIX_CELL_PROMPT = """You are writing a compliance-matrix cell response for a Technical \
Proposal responding to a government tender.

REQUIREMENT (from compliance matrix row):
{requirement}

TENT TYPE: {tent_label}
DOCUMENTATION GUIDANCE (evaluator expectations):
{doc_guidance}

RETRIEVED TECHNICAL EVIDENCE:
{evidence}

Respond in this exact format with no deviations:
YES_NO: [YES or NO]
CROSS_REFERENCE: [a short descriptive title for this response]
REMARK: [max 80 words — cite specific evidence, be precise, flag limitations honestly]
PROPOSAL_TEXT:
[100-200 words grounded in the retrieved evidence. Do not invent specifications or certifications.]
"""


def _load_matrix_requirements(matrix_path):
    df = pd.read_excel(matrix_path, header=None)
    requirements = []
    for idx, row in df.iterrows():
        try:
            req_num = int(float(str(row.iloc[COL_NO])))
        except (TypeError, ValueError):
            continue
        req_text = str(row.iloc[COL_REQ]).strip()
        if req_text and req_text not in ("nan", "NaN"):
            requirements.append({
                "row_idx": idx, "num": req_num, "requirement": req_text,
                "sow_ref": str(row.iloc[COL_SOW_REF]).strip(),
                "doc_guidance": str(row.iloc[COL_DOC_GUIDE]).strip(),
            })
    return requirements


def _generate_matrix_cell(requirement, doc_guidance, tent_label, evidence_chunks):
    prompt = _MATRIX_CELL_PROMPT.format(
        requirement=requirement, tent_label=tent_label, doc_guidance=doc_guidance,
        evidence=_evidence_block(evidence_chunks))
    import anthropic
    client = anthropic.Anthropic()
    message = client.messages.create(model=CLAUDE_MODEL, max_tokens=800,
                                      messages=[{"role": "user", "content": prompt}])
    return _parse_matrix_cell_response(message.content[0].text)


def _parse_matrix_cell_response(raw):
    result = {"yes_no": "?", "cross_ref": "", "remark": "", "proposal_text": ""}
    in_proposal, proposal_lines = False, []
    for line in raw.split("\n"):
        if line.startswith("YES_NO:"):
            result["yes_no"] = line.replace("YES_NO:", "").strip()
        elif line.startswith("CROSS_REFERENCE:"):
            result["cross_ref"] = line.replace("CROSS_REFERENCE:", "").strip()
        elif line.startswith("REMARK:"):
            result["remark"] = line.replace("REMARK:", "").strip()
        elif line.startswith("PROPOSAL_TEXT:"):
            in_proposal = True
        elif in_proposal:
            proposal_lines.append(line)
    result["proposal_text"] = "\n".join(proposal_lines).strip()
    return result


def fill_compliance_matrix(tenant_id, pub_number, matrix_path, out_path):
    """Requires ANTHROPIC_API_KEY for any TO_BE_LINKED/COMPLETE cell (raises
    via anthropic's own client if unset — unlike the SOW path, there's no
    graceful partial-fill mode here since a matrix with unset cells isn't a
    useful deliverable; call only when a key is configured).
    """
    requirements = _load_matrix_requirements(matrix_path)
    wb = openpyxl.load_workbook(matrix_path)
    ws = wb.active

    green = PatternFill("solid", fgColor="C6EFCE")
    red = PatternFill("solid", fgColor="FFC7CE")
    amber = PatternFill("solid", fgColor="FFEB9C")
    blue = PatternFill("solid", fgColor="DDEBF7")

    for req in requirements:
        row_num = req["row_idx"] + 1
        for tent in TENT_SIZES:
            query = f"{req['requirement']} {req['doc_guidance']}"
            tech_chunks = retrieve_evidence(tenant_id, pub_number, query, roles=["tech"])
            status = _gap_status(tech_chunks)

            if status == "completed":
                fill, yesno_val = amber, "TO BE COMPLETED"
                remark_val = "No technical documentation found. To be completed before submission."
                ref_val = f"Section {req['num']}"
            else:
                parsed = _generate_matrix_cell(req["requirement"], req["doc_guidance"],
                                                tent["label"], tech_chunks)
                if status == "linked":
                    best_doc = max(tech_chunks, key=lambda c: c["similarity"])["source"]
                    fill = blue
                    remark_val = f"{parsed['remark']} | Document: {best_doc} — to be linked"
                else:
                    fill = green if parsed["yes_no"] == "YES" else red
                    remark_val = parsed["remark"]
                yesno_val = parsed["yes_no"]
                ref_val = parsed["cross_ref"]

            yn_cell = ws.cell(row=row_num, column=tent["col_yesno"] + 1)
            ref_cell = ws.cell(row=row_num, column=tent["col_ref"] + 1)
            rem_cell = ws.cell(row=row_num, column=tent["col_remark"] + 1)
            yn_cell.value, yn_cell.fill = yesno_val, fill
            ref_cell.value = ref_val
            rem_cell.value = remark_val
            if status in ("completed", "linked"):
                rem_cell.font = XLFont(color="9C0006", italic=True)

    wb.save(out_path)
    return out_path
